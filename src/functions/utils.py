import io
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
import os
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import sys

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)
FONT_PATH = os.path.join(BASE_DIR, "assets", "DejaVuSans.ttf")

pdfmetrics.registerFont(TTFont('DejaVu', FONT_PATH))


def export_excel(record):
    """Trả về bytes của file Excel"""
    df = pd.DataFrame([record])
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False, engine='openpyxl')
    buffer.seek(0)
    return buffer.getvalue()  # Trả về bytes

def export_word(record):
    """Trả về bytes của file Word"""
    doc = Document()
    doc.add_heading('Báo cáo chẩn đoán', 0)
    
    for key, value in record.items():
        doc.add_paragraph(f"{key}: {value}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Đăng ký 1 lần (có thể đặt ở đầu file hoặc __init__)

def export_pdf(record):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    
    c.setFont('DejaVu', 12)  # Dùng font đã đăng ký
    
    y = 800
    c.drawString(100, y, "BÁO CÁO CHẨN ĐOÁN")
    y -= 30
    
    for key, value in record.items():
        c.drawString(100, y, f"{key}: {value}")
        y -= 20
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()